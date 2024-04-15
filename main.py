# Importing the Required Libraries
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import pandas
import docx
import datetime
from docx.shared import Pt
from docx2pdf import convert
from placeholders import MY_EMAIL,PASSWORD,SENDERS_NAME, ADDRESS_LINE_1,ADDRESS_LINE_2,ADDRESS_LINE_3,ADDRESS_LINE_4,SCHOOL_NAME,BANK_NAME,SORT_CODE,ACCOUNT_NUMBER,FULL_NAME

# Finding out the Date
today = datetime.date.today().strftime("%d/%m/%y")
month = datetime.date.today().month

# Working out the Season
season = ""

if 1 <= month <= 2 or month == 12:
    season = "Winter"
elif 3 <= month <= 5:
    season = "Spring"
elif 6 <= month <= 8:
    season = "Summer"
else:
    season = "Autumn"

# Creating the Invoice
def create_invoice(index_number):
    # Creating the word document from invoice template
    doc = docx.Document(f"Input/invoice_template.docx")

    # Set font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Helvetica'

    #Replacing Address Placeholders
    address = doc.paragraphs[0]
    address.text = f"{SENDERS_NAME}\n{ADDRESS_LINE_1}\n{ADDRESS_LINE_2}\n{ADDRESS_LINE_3}\n{ADDRESS_LINE_4}"
    # run = address.runs[0]

    # Replacing Date Placeholder
    date_para = doc.paragraphs[2]
    date_para.text = today
    # run = date_para.runs[0]

    # Replacing Placeholders in Title
    title = doc.paragraphs[4]
    title.text = f"Invoice for the {student_dict['instrument'][index_number].title()} Lessons during the {season} Term"
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(14)

    location = doc.paragraphs[5]
    location.text = f"At: {SCHOOL_NAME}"
    run = location.runs[0]
    run.bold = True
    run.font.size = Pt(14)

    # Replacing Placeholders in Introduction
    first_line = doc.paragraphs[7]
    first_line.text = f"Dear the parent/guardian of {student_dict['name'][index_number]},"

    # Replacing Placeholders in Price of Lessons and Total
    fee_line = doc.paragraphs[11]
    lesson_price = "{:.2f}".format(student_dict['price'][index_number])
    fee_line.text = f"Fee agreed for the teaching at £{lesson_price} per {student_dict['lesson_length'][index_number]}. "

    total_line = doc.paragraphs[14]
    TOTAL = "{:.2f}".format(10 * student_dict['price'][index_number])
    total_line.text = f"Total: £{TOTAL}"

    #Replacing Bank Details Placeholders
    bank_details = doc.paragraphs[18]
    bank_details.text = f"{BANK_NAME}\n{FULL_NAME}\nSort Code: {SORT_CODE}\nAccount Number:{ACCOUNT_NUMBER}"

    # Replacing Name Placeholder
    sign_off = doc.paragraphs[24]
    sign_off.text = SENDERS_NAME

    # Saving edited document
    doc.save(f"Output/Invoices/{student_dict['name'][index_number]} {season} Term Invoice.docx")

# Creating the Email Object
def send_email(index_number):
    with open(f"Input/email.txt") as content:
        # Email subject
        SUBJECT = f"Invoice for {season} Term Peripatetic Music Lessons"
        # Reading email body template
        email_contents = content.read()
        # Finding Student Information
        student_email = student_dict["email"][index_number]
        student_name = student_dict["name"][index_number].title()
        student_form = student_dict["form"][index_number]
        student_instrument = student_dict["instrument"][index_number]
        # Replacing Placeholders with Student Information
        body = email_contents.replace("[NAME]", student_name)
        body = body.replace("[FIRST_NAME]", student_name.split(" ")[0])
        body = body.replace("[FORM]", student_form)
        body = body.replace("[INSTRUMENT]", student_instrument)
        body = body.replace("[SENDERS_NAME]", SENDERS_NAME.split(" ")[0])

    # Creating email object
    message = MIMEMultipart()
    message["From"] = MY_EMAIL
    message["To"] = student_email
    message["Subject"] = SUBJECT

    # Add the body to the email
    message.attach(MIMEText(body, "plain"))

    # Attaching the File
    attachment_path = f"Output/PDFs/{student_dict['name'][index]} {season} term Invoice.pdf"

    # Attach the file using the MIMEBase class
    attachment = open(attachment_path, "rb")
    payload = MIMEBase("application", "octet-stream")
    payload.set_payload((attachment).read())
    encoders.encode_base64(payload)
    payload.add_header(
        "Content-Disposition", f'attachment; filename= "{attachment_path.split("/")[-1]}"'
    )
    message.attach(payload)

    # Establishing the SMTP Connection
    smtp_server = "smtp.gmail.com"
    smtp_port = 587

    with smtplib.SMTP(smtp_server, smtp_port) as server:
        server.starttls()
        server.login(MY_EMAIL, PASSWORD)
        server.send_message(msg=message)

# Creating Dictionary from Dataframe
df = pandas.read_csv("Input/names.csv")
student_dict = df.to_dict(orient="list")

# Creating Invoice, Converting Invoice to PDF, and Emailing Invoice For Each Student in Dictionary
for index in range(len(student_dict['name'])):
    create_invoice(index)
    convert(f"Output/Invoices/{student_dict['name'][index]} {season} Term Invoice.docx", f"Output/PDFs/{student_dict['name'][index]} {season} term invoice.pdf")
    send_email(index)




